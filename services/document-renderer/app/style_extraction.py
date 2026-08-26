"""
Parses a .docx template into a StyleCatalog: the set of paragraph styles,
numbering formats, fonts, and tables actually used in the document body.

This catalog is what gets shown to Claude ("these are the only styles/numbering
you may reference"), and it's also used by the renderer to know how to attach
numbering to a node that only specifies a style_name.
"""
from __future__ import annotations

from collections import OrderedDict
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table

from app.schemas import (
    FontInfo,
    NumberingInfo,
    StyleCatalog,
    StyleCatalogEntry,
    TableCatalogEntry,
)


def _build_numbering_format_maps(document: Document):
    """Returns (num_to_abstract: {numId(str): abstractNumId(str)},
    abstract_formats: {abstractNumId(str): {ilvl(int): fmt(str)}})."""
    try:
        numbering_part = document.part.numbering_part
    except Exception:
        return {}, {}

    xml = numbering_part.element
    num_to_abstract: dict[str, str] = {}
    for num in xml.findall(qn("w:num")):
        num_id = num.get(qn("w:numId"))
        abstract_el = num.find(qn("w:abstractNumId"))
        if num_id is not None and abstract_el is not None:
            num_to_abstract[num_id] = abstract_el.get(qn("w:val"))

    abstract_formats: dict[str, dict[int, str]] = {}
    for abstract_num in xml.findall(qn("w:abstractNum")):
        abstract_id = abstract_num.get(qn("w:abstractNumId"))
        lvl_map: dict[int, str] = {}
        for lvl in abstract_num.findall(qn("w:lvl")):
            ilvl_raw = lvl.get(qn("w:ilvl"))
            fmt_el = lvl.find(qn("w:numFmt"))
            if ilvl_raw is not None and fmt_el is not None:
                lvl_map[int(ilvl_raw)] = fmt_el.get(qn("w:val"))
        if abstract_id is not None:
            abstract_formats[abstract_id] = lvl_map

    return num_to_abstract, abstract_formats


def _direct_num_pr(paragraph):
    """numPr set directly on the paragraph (not via its style)."""
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return None
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    num_id_el = num_pr.find(qn("w:numId"))
    ilvl_el = num_pr.find(qn("w:ilvl"))
    if num_id_el is None:
        return None
    num_id = num_id_el.get(qn("w:val"))
    ilvl = ilvl_el.get(qn("w:val")) if ilvl_el is not None else "0"
    return num_id, int(ilvl)


def _style_chain_num_pr(paragraph):
    """Walk up paragraph.style -> base_style looking for numPr defined on the style itself."""
    style = paragraph.style
    seen = set()
    while style is not None and id(style) not in seen:
        seen.add(id(style))
        style_el = style.element
        p_pr = style_el.find(qn("w:pPr"))
        if p_pr is not None:
            num_pr = p_pr.find(qn("w:numPr"))
            if num_pr is not None:
                num_id_el = num_pr.find(qn("w:numId"))
                ilvl_el = num_pr.find(qn("w:ilvl"))
                if num_id_el is not None:
                    num_id = num_id_el.get(qn("w:val"))
                    ilvl = ilvl_el.get(qn("w:val")) if ilvl_el is not None else "0"
                    return num_id, int(ilvl)
        style = getattr(style, "base_style", None)
    return None


def _effective_numbering(paragraph, num_to_abstract, abstract_formats) -> NumberingInfo | None:
    found = _direct_num_pr(paragraph) or _style_chain_num_pr(paragraph)
    if not found:
        return None
    num_id, ilvl = found
    abstract_id = num_to_abstract.get(num_id)
    fmt = abstract_formats.get(abstract_id, {}).get(ilvl) if abstract_id else None
    return NumberingInfo(numId=int(num_id), ilvl=ilvl, format=fmt)


def _font_info(paragraph) -> FontInfo | None:
    runs = [r for r in paragraph.runs if r.text.strip()]
    run = runs[0] if runs else None

    def _first_non_none(getter):
        if run is not None:
            value = getter(run.font)
            if value is not None:
                return value
        style = paragraph.style
        seen = set()
        while style is not None and id(style) not in seen:
            seen.add(id(style))
            value = getter(style.font)
            if value is not None:
                return value
            style = getattr(style, "base_style", None)
        return None

    name = _first_non_none(lambda f: f.name)
    size = _first_non_none(lambda f: f.size)
    bold = _first_non_none(lambda f: f.bold)
    italic = _first_non_none(lambda f: f.italic)
    return FontInfo(name=name, size_pt=size.pt if size else None, bold=bold, italic=italic)


def _table_catalog_entry(table: Table) -> TableCatalogEntry:
    example_rows = []
    for row in table.rows[:2]:
        example_rows.append(" | ".join(cell.text.strip() for cell in row.cells))
    return TableCatalogEntry(
        style_name=table.style.name if table.style else None,
        row_count=len(table.rows),
        col_count=len(table.columns),
        example_row_texts=example_rows,
    )


def extract_style_catalog(docx_bytes: bytes) -> StyleCatalog:
    document = Document(BytesIO(docx_bytes))
    num_to_abstract, abstract_formats = _build_numbering_format_maps(document)

    seen_styles: "OrderedDict[str, StyleCatalogEntry]" = OrderedDict()
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else "Normal"
        if style_name in seen_styles:
            seen_styles[style_name].usage_count += 1
            continue
        numbering = _effective_numbering(paragraph, num_to_abstract, abstract_formats)
        seen_styles[style_name] = StyleCatalogEntry(
            style_name=style_name,
            style_id=paragraph.style.style_id if paragraph.style else "Normal",
            usage_count=1,
            example_text=text[:120],
            numbering=numbering,
            font=_font_info(paragraph),
        )

    tables = [_table_catalog_entry(t) for t in document.tables]

    return StyleCatalog(paragraph_styles=list(seen_styles.values()), tables=tables)
