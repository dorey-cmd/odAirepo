"""
Renders a new .docx by starting from a copy of the original template (so
styles.xml, numbering.xml, theme, and section/page setup are preserved
byte-for-byte), wiping only the body content, and rebuilding it from an
AI-produced ContentTree that references style/numbering IDs from the
template's own StyleCatalog.
"""
from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.schemas import ContentNode, ContentTree

_ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# A body paragraph the model already numbered as literal text ("6.3.1 ...",
# "(א) ...", "23. ...") should never ALSO get Word's own numId/ilvl auto-
# numbering - that renders both, producing visibly doubled numbers. This is
# a deterministic backstop: even if the prompt guidance is ignored, the
# renderer itself won't let both apply to the same paragraph.
_LEADING_NUMBER_RE = re.compile(r"^\s*(\(?[0-9]+([.\-][0-9]+)*\)?[.\)]?|\(?[א-ת]\)?[.\)])\s")

# Legal templates typed directly in Word are commonly typed with the
# style's own space-after left at zero, which reads as a solid wall of text
# once assembled from many short paragraphs. Guaranteeing a minimum here -
# rather than asking the model to somehow know to add spacing - is a system-
# level formatting guarantee, not a per-document AI judgment call.
_MIN_SPACE_AFTER_PT = 8


def _looks_pre_numbered(text: str | None) -> bool:
    return bool(text and _LEADING_NUMBER_RE.match(text.strip()))


def _effective_space_after_pt(paragraph) -> float | None:
    style = paragraph.style
    seen = set()
    while style is not None and id(style) not in seen:
        seen.add(id(style))
        space_after = style.paragraph_format.space_after
        if space_after is not None:
            return space_after.pt
        style = getattr(style, "base_style", None)
    return None


def _ensure_min_spacing(paragraph) -> None:
    direct = paragraph.paragraph_format.space_after
    if direct is not None and direct.pt >= _MIN_SPACE_AFTER_PT:
        return
    inherited = _effective_space_after_pt(paragraph)
    if inherited is not None and inherited >= _MIN_SPACE_AFTER_PT:
        return
    paragraph.paragraph_format.space_after = Pt(_MIN_SPACE_AFTER_PT)


def _clear_body_keep_section(document: Document) -> None:
    body = document.element.body
    section_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not section_pr:
            body.remove(child)


def _set_num_pr(paragraph, num_id: int, ilvl: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.insert(0, num_pr)


def _known_paragraph_style_names(document: Document) -> set[str]:
    return {s.name for s in document.styles if s.type is not None and s.name}


def _apply_flag(document: Document, paragraph, flag_text: str) -> None:
    for run in paragraph.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(_MIN_SPACE_AFTER_PT)
    run = note.add_run(f"⚑ לתשומת לב עורך/ת הדין: {flag_text}")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x4D, 0x00)


def _add_paragraph(document: Document, node: ContentNode, known_styles: set[str]):
    style_name = node.style_name if node.style_name in known_styles else None
    paragraph = document.add_paragraph(node.text or "", style=style_name)
    if node.numId is not None and not _looks_pre_numbered(node.text):
        _set_num_pr(paragraph, node.numId, node.ilvl or 0)
    else:
        # The paragraph's STYLE itself (e.g. a heading style) can carry its
        # own baked-in numPr in the template's styles.xml, independent of
        # anything set on this specific paragraph. Left alone, Word silently
        # applies that inherited auto-number on top of whatever literal
        # number the model already typed into the text (e.g. "1.1 ...") -
        # producing a visibly doubled number that a node-level check for
        # "does this node have both numId and a leading number" can never
        # catch, since the paragraph's OWN numId is genuinely None. Setting
        # numId=0 is OOXML's standard way to explicitly override/disable an
        # inherited style numbering for one paragraph.
        _set_num_pr(paragraph, 0, 0)
    if node.alignment:
        paragraph.alignment = _ALIGNMENT_MAP.get(node.alignment)
    _ensure_min_spacing(paragraph)
    if node.flag:
        _apply_flag(document, paragraph, node.flag)
    return paragraph


def _add_table(document: Document, node: ContentNode, known_styles: set[str]):
    rows = node.rows or [[""]]
    n_cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=n_cols)
    if node.style_name:
        try:
            table.style = node.style_name
        except KeyError:
            pass
    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            table.rows[r_idx].cells[c_idx].text = row[c_idx] if c_idx < len(row) else ""
    return table


def render_document(template_bytes: bytes, content_tree: ContentTree) -> bytes:
    document = Document(BytesIO(template_bytes))
    known_styles = _known_paragraph_style_names(document)
    _clear_body_keep_section(document)

    for node in content_tree.nodes:
        if node.type == "table":
            _add_table(document, node, known_styles)
        else:
            _add_paragraph(document, node, known_styles)

    out = BytesIO()
    document.save(out)
    return out.getvalue()
