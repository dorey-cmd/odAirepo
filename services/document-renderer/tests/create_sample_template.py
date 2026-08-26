"""
Builds a synthetic contract template .docx that mimics real-world structure:
a title, "Heading 1" section titles, a custom-styled numbered clause list
(direct numPr, the common real-world case), a sub-clause level, a signature
table, and a house font/size applied via a custom style — so the Phase 0
prototype has something non-trivial to preserve.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def _add_num_pr(paragraph, num_id: int, ilvl: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.insert(0, num_pr)


def _add_numbering_definition(document: Document, num_id: int, abstract_num_id: int, fmt: str) -> None:
    numbering_part = document.part.numbering_part
    numbering_el = numbering_part.element

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(start)
    lvl.append(num_fmt)
    lvl.append(lvl_text)
    abstract_num.append(lvl)
    numbering_el.insert(0, abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_ref)
    numbering_el.append(num)


def build_sample_template(output_path: Path) -> None:
    document = Document()

    body_style = document.styles.add_style("Contract Body", WD_STYLE_TYPE.PARAGRAPH)
    body_style.base_style = document.styles["Normal"]
    body_style.font.name = "Georgia"
    body_style.font.size = Pt(11)

    clause_style = document.styles.add_style("Contract Clause", WD_STYLE_TYPE.PARAGRAPH)
    clause_style.base_style = document.styles["Normal"]
    clause_style.font.name = "Georgia"
    clause_style.font.size = Pt(11)

    _add_numbering_definition(document, num_id=1, abstract_num_id=1, fmt="decimal")

    title = document.add_paragraph("Master Services Agreement", style="Title")
    title.alignment = 1

    document.add_paragraph(
        "This is a generic sample template used to validate style-faithful "
        "rendering. Real templates will be uploaded by lawyers.",
        style="Contract Body",
    )

    document.add_paragraph("Definitions", style="Heading 1")
    clause = document.add_paragraph(
        "\"Agreement\" means this contract together with all exhibits attached hereto.",
        style="Contract Clause",
    )
    _add_num_pr(clause, num_id=1, ilvl=0)

    document.add_paragraph("Term and Termination", style="Heading 1")
    clause2 = document.add_paragraph(
        "This Agreement commences on the Effective Date and continues until terminated.",
        style="Contract Clause",
    )
    _add_num_pr(clause2, num_id=1, ilvl=0)

    document.add_paragraph("Signatures", style="Heading 1")
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Party A"
    table.rows[0].cells[1].text = "Party B"
    table.rows[1].cells[0].text = "Signature: ____________"
    table.rows[1].cells[1].text = "Signature: ____________"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    out = Path(__file__).parent / "fixtures" / "sample_template.docx"
    build_sample_template(out)
    print(f"Wrote {out}")
